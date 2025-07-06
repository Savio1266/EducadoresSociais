# app.py
from flask import Flask, render_template, request, redirect, url_for, session, flash, send_file
from flask_sqlalchemy import SQLAlchemy
from werkzeug.security import generate_password_hash, check_password_hash
from datetime import datetime
import re
from unicodedata import category
from docx import Document  # Importa a biblioteca python-docx
from docx.shared import Inches  # Para ajustar margens se necessário
import io

app = Flask(__name__)
app.secret_key = 'segredo_seguro'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///diario_cef07.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)


# -------------------- MODELOS -------------------- #

class Usuario(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    nome_usuario = db.Column(db.String(50), unique=True, nullable=False)
    senha_hash = db.Column(db.String(200), nullable=False)
    tipo_acesso = db.Column(db.String(20), nullable=False)
    status = db.Column(db.String(20), default='pendente')

    def verificar_senha(self, senha):
        return check_password_hash(self.senha_hash, senha)


class Estudante(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    nome = db.Column(db.String(100), nullable=False)
    turma = db.Column(db.String(20), nullable=False)
    turno = db.Column(db.String(20), nullable=False)


class EducadorEstudante(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    educador_id = db.Column(db.Integer, db.ForeignKey('usuario.id'))
    estudante_id = db.Column(db.Integer, db.ForeignKey('estudante.id'))


class Registro(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    estudante_id = db.Column(db.Integer, db.ForeignKey('estudante.id'))
    educador_id = db.Column(db.Integer, db.ForeignKey('usuario.id'))
    data = db.Column(db.String(20))
    disciplina = db.Column(db.String(50))
    qtd_aulas = db.Column(db.Integer)
    texto = db.Column(db.Text)


# -------------------- FUNÇÕES AUXILIARES -------------------- #

def usuario_logado():
    """Verifica se um usuário está logado."""
    return session.get('usuario_id') is not None


def tipo_usuario():
    """Retorna o tipo de acesso do usuário logado."""
    return session.get('tipo_acesso')


def sanitize_text_for_docx(text):
    """
    Sanitiza o texto para inclusão em documentos DOCX.
    A biblioteca python-docx geralmente lida bem com Unicode,
    então esta função é mais para garantir quebra de linha em
    palavras muito longas e remover caracteres de controle problemáticos.
    """
    # Remove caracteres de controle que podem causar problemas (exceto quebras de linha)
    text = ''.join(c for c in text if category(c)[0] != 'C' or c in ('\n', '\r', '\t'))
    # Adiciona quebras de linha suaves em palavras muito longas para melhorar a formatação em Word
    text = re.sub(r'(\S{50})(?=\S)', r'\1 ', text)  # Quebra palavras com mais de 50 caracteres
    return text.strip()


# -------------------- ROTAS GERAIS -------------------- #

@app.route('/')
def index():
    """Rota da página inicial."""
    return render_template('index.html')


@app.route('/login/<tipo>', methods=['GET', 'POST'])
def login(tipo):
    """Rota de login para diferentes tipos de usuário."""
    if request.method == 'POST':
        nome = request.form['nome_usuario']
        senha = request.form['senha']
        usuario = Usuario.query.filter_by(nome_usuario=nome, tipo_acesso=tipo).first()
        if usuario and usuario.status == 'aprovado' and usuario.verificar_senha(senha):
            session['usuario_id'] = usuario.id
            session['usuario_nome'] = usuario.nome_usuario
            session['tipo_acesso'] = usuario.tipo_acesso
            if tipo in ['Moderador', 'Direcao', 'Coordenacao']:
                return redirect(url_for('painel_moderador'))
            elif tipo == 'Educador Social':
                return redirect(url_for('painel_educador_social'))
            else:
                return redirect(url_for('painel_educador'))
        else:
            flash('Credenciais inválidas ou cadastro pendente.', 'danger')
    return render_template('login.html', tipo=tipo)


@app.route('/cadastro/<tipo>', methods=['GET', 'POST'])
def cadastro(tipo):
    """Rota de cadastro para novos usuários."""
    if request.method == 'POST':
        nome = request.form['nome_usuario']
        senha = request.form['senha']
        if Usuario.query.filter_by(nome_usuario=nome).first():
            flash('Nome de usuário já existe.', 'warning')
            return redirect(url_for('cadastro', tipo=tipo))
        senha_hash = generate_password_hash(senha)
        novo = Usuario(nome_usuario=nome, senha_hash=senha_hash, tipo_acesso=tipo)
        db.session.add(novo)
        db.session.commit()
        flash('Cadastro enviado para aprovação do moderador.', 'success')
        return redirect(url_for('login', tipo=tipo))
    return render_template('cadastro.html', tipo=tipo)


@app.route('/logout')
def logout():
    """Rota para encerrar a sessão do usuário."""
    session.clear()
    return redirect(url_for('index'))


# -------------------- PAINEL MODERADOR -------------------- #

@app.route('/painel_moderador')
def painel_moderador():
    """Painel de controle para moderadores, diretores e coordenadores."""
    if tipo_usuario() not in ['Moderador', 'Direcao', 'Coordenacao']:
        flash('Acesso não autorizado.', 'danger')
        return redirect(url_for('index'))
    pendentes = Usuario.query.filter_by(status='pendente').all()
    usuarios = Usuario.query.filter_by(status='aprovado').all()
    estudantes = Estudante.query.all()
    educadores = Usuario.query.filter(Usuario.tipo_acesso.in_(['Educador', 'Educador Social']),
                                      Usuario.status == 'aprovado').all()
    # O template dashboard.html para moderador está na subpasta 'moderador'
    return render_template('moderador/dashboard.html', pendentes=pendentes, usuarios=usuarios, estudantes=estudantes,
                           educadores=educadores)


@app.route('/aprovar/<int:id>')
def aprovar(id):
    """Aprova o cadastro de um usuário pendente."""
    if tipo_usuario() not in ['Moderador', 'Direcao', 'Coordenacao']:
        flash('Acesso não autorizado.', 'danger')
        return redirect(url_for('index'))
    usuario = Usuario.query.get(id)
    if usuario:
        usuario.status = 'aprovado'
        db.session.commit()
        flash(f'Usuário {usuario.nome_usuario} aprovado.', 'success')
    return redirect(url_for('painel_moderador'))


@app.route('/recusar/<int:id>')
def recusar(id):
    """Recusa e remove o cadastro de um usuário pendente."""
    if tipo_usuario() not in ['Moderador', 'Direcao', 'Coordenacao']:
        flash('Acesso não autorizado.', 'danger')
        return redirect(url_for('index'))
    usuario = Usuario.query.get(id)
    if usuario:
        db.session.delete(usuario)
        db.session.commit()
        flash(f'Usuário {usuario.nome_usuario} recusado e removido.', 'info')
    return redirect(url_for('painel_moderador'))


@app.route('/excluir_usuario/<int:id>')
def excluir_usuario(id):
    """Exclui um usuário ativo e seus dados associados."""
    if tipo_usuario() not in ['Moderador', 'Direcao', 'Coordenacao']:
        flash('Acesso não autorizado.', 'danger')
        return redirect(url_for('index'))

    usuario = Usuario.query.get(id)
    if not usuario:
        flash('Usuário não encontrado.', 'danger')
        return redirect(url_for('painel_moderador'))

    if usuario.id == session.get('usuario_id'):
        flash('Você não pode excluir seu próprio cadastro de moderador.', 'warning')
        return redirect(url_for('painel_moderador'))

    if usuario.tipo_acesso in ['Educador', 'Educador Social']:
        EducadorEstudante.query.filter_by(educador_id=usuario.id).delete()
        Registro.query.filter_by(educador_id=usuario.id).delete()
        db.session.commit()

    db.session.delete(usuario)
    db.session.commit()
    flash(f'Usuário {usuario.nome_usuario} e todos os seus dados associados foram excluídos.', 'success')
    return redirect(url_for('painel_moderador'))


@app.route('/cadastrar_estudante', methods=['POST'])
def cadastrar_estudante():
    """Cadastra um novo estudante no sistema."""
    if tipo_usuario() not in ['Moderador', 'Direcao', 'Coordenacao']:
        flash('Acesso não autorizado.', 'danger')
        return redirect(url_for('index'))
    nome = request.form['nome']
    turma = request.form['turma']
    turno = request.form['turno']
    novo = Estudante(nome=nome, turma=turma, turno=turno)
    db.session.add(novo)
    db.session.commit()
    flash(f'Estudante {nome} cadastrado com sucesso!', 'success')
    return redirect(url_for('painel_moderador'))


@app.route('/excluir_estudante/<int:estudante_id>')
def excluir_estudante(estudante_id):
    """Exclui um estudante e seus dados associados."""
    if tipo_usuario() not in ['Moderador', 'Direcao', 'Coordenacao']:
        flash('Acesso não autorizado.', 'danger')
        return redirect(url_for('index'))
    estudante = Estudante.query.get(estudante_id)
    if estudante:
        EducadorEstudante.query.filter_by(estudante_id=estudante_id).delete()
        Registro.query.filter_by(estudante_id=estudante_id).delete()
        db.session.delete(estudante)
        db.session.commit()
        flash(f'Estudante {estudante.nome} e seus dados foram excluídos.', 'success')
    return redirect(url_for('painel_moderador'))


@app.route('/vinculos/<int:educador_id>')
def vinculos(educador_id):
    """Exibe os estudantes vinculados a um educador específico."""
    if tipo_usuario() not in ['Moderador', 'Direcao', 'Coordenacao']:
        flash('Acesso não autorizado.', 'danger')
        return redirect(url_for('index'))
    estudantes = db.session.query(Estudante).join(EducadorEstudante).filter(
        EducadorEstudante.educador_id == educador_id).all()
    educador = Usuario.query.get(educador_id)
    # CORREÇÃO AQUI: O template vinculos.html provavelmente está na subpasta 'moderador'
    return render_template('moderador/vinculos.html', estudantes=estudantes,
                           educador=educador)


# -------------------- PAINEL EDUCADORES -------------------- #

@app.route('/painel_educador')
def painel_educador():
    """Painel de controle para educadores."""
    if tipo_usuario() not in ['Educador', 'Educador Social']:
        flash('Acesso não autorizado.', 'danger')
        return redirect(url_for('index'))
    meus_estudantes = db.session.query(Estudante).join(EducadorEstudante).filter(
        EducadorEstudante.educador_id == session['usuario_id']).all()
    todos = Estudante.query.all()
    vinculados_ids = [e.id for e in meus_estudantes]
    nao_vinculados = [e for e in todos if e.id not in vinculados_ids]
    return render_template('educador/dashboard.html', estudantes=meus_estudantes, disponiveis=nao_vinculados)


@app.route('/painel_educador_social')
def painel_educador_social():
    """Redireciona para o painel do educador."""
    return redirect(url_for('painel_educador'))


@app.route('/vincular_estudante', methods=['POST'])
def vincular_estudante():
    """Vincula um estudante ao educador logado."""
    if tipo_usuario() not in ['Educador', 'Educador Social']:
        flash('Acesso não autorizado.', 'danger')
        return redirect(url_for('index'))
    estudante_id = request.form['estudante_id']
    ja_vinculado = EducadorEstudante.query.filter_by(educador_id=session['usuario_id'],
                                                     estudante_id=estudante_id).first()
    if not ja_vinculado:
        vinculo = EducadorEstudante(educador_id=session['usuario_id'], estudante_id=estudante_id)
        db.session.add(vinculo)
        db.session.commit()
        flash(f'Estudante vinculado com sucesso!', 'success')
    else:
        flash(f'Estudante já está vinculado.', 'info')
    return redirect(url_for('painel_educador'))


@app.route('/registrar_aula', methods=['GET', 'POST'])
@app.route('/registrar_aula/<int:estudante_id>', methods=['GET', 'POST'])
def registrar_aula(estudante_id=None):
    """Registra uma nova aula para um ou mais estudantes."""
    if tipo_usuario() not in ['Educador', 'Educador Social']:
        flash('Acesso não autorizado.', 'danger')
        return redirect(url_for('index'))

    vinculados = db.session.query(Estudante).join(EducadorEstudante).filter(
        EducadorEstudante.educador_id == session['usuario_id']
    ).all()

    estudantes_a_selecionar = []
    if estudante_id:
        estudante_unico = Estudante.query.get(estudante_id)
        if estudante_unico and any(e.id == estudante_id for e in vinculados):
            estudantes_a_selecionar = [estudante_id]
        else:
            flash(f'Estudante selecionado não está vinculado a você ou não existe.', 'warning')
            estudante_id = None

    if request.method == 'POST':
        estudantes_ids = request.form.getlist('estudantes')
        data = request.form['data']
        disciplina = request.form['disciplina']
        qtd_aulas = int(request.form['qtd_aulas'])
        texto = request.form['texto']

        if not estudantes_ids:
            flash('Selecione pelo menos um estudante para registrar a aula.', 'warning')
            return render_template('educador/registrar_aula.html', estudantes=vinculados,
                                   estudantes_a_selecionar=estudantes_a_selecionar, now=datetime.now)

        for estudante_id_reg in estudantes_ids:
            registro = Registro(
                estudante_id=estudante_id_reg,
                educador_id=session['usuario_id'],
                data=data,
                disciplina=disciplina,
                qtd_aulas=qtd_aulas,
                texto=texto
            )
            db.session.add(registro)
        db.session.commit()
        flash('Registro de aula(s) salvo(s) com sucesso!', 'success')

        return redirect(url_for('ver_registros'))

    return render_template('educador/registrar_aula.html', estudantes=vinculados,
                           estudantes_a_selecionar=estudantes_a_selecionar, now=datetime.now)


@app.route('/editar_registro/<int:registro_id>', methods=['GET', 'POST'])
def editar_registro(registro_id):
    """Edita um registro de aula existente."""
    registro = Registro.query.get_or_404(registro_id)

    if not (registro.educador_id == session.get('usuario_id') or
            tipo_usuario() in ['Moderador', 'Direcao', 'Coordenacao']):
        flash('Você não tem permissão para editar este registro.', 'danger')
        return redirect(url_for('painel_educador'))

    if request.method == 'POST':
        registro.data = request.form['data']
        registro.disciplina = request.form['disciplina']
        registro.qtd_aulas = int(request.form['qtd_aulas'])
        registro.texto = request.form['texto']
        db.session.commit()
        flash('Registro atualizado com sucesso!', 'success')

        from_student_history = request.args.get('from_student_history', 'false').lower() == 'true'
        if from_student_history:
            return redirect(url_for('ver_registros', estudante_id=registro.estudante_id))
        else:
            return redirect(url_for('ver_registros'))

    return render_template('educador/editar_registro.html', registro=registro)


@app.route('/ver_registros')
@app.route('/ver_registros/<int:estudante_id>')
def ver_registros(estudante_id=None):
    """Exibe os registros de aula do educador logado ou de um estudante específico."""
    if tipo_usuario() not in ['Educador', 'Educador Social', 'Moderador', 'Direcao', 'Coordenacao']:
        flash('Acesso não autorizado.', 'danger')
        return redirect(url_for('index'))

    base_query = Registro.query
    if tipo_usuario() in ['Educador', 'Educador Social']:
        base_query = base_query.filter_by(educador_id=session['usuario_id'])

    estudante_info = None

    if estudante_id:
        estudante_info = Estudante.query.get_or_404(estudante_id)
        registros_para_exibir = base_query.filter_by(estudante_id=estudante_id).order_by(Registro.data.desc()).all()
        return render_template('educador/ver_registros.html',
                               estudante=estudante_info,
                               registros=registros_para_exibir,
                               is_student_history=True)
    else:
        all_registros = base_query.order_by(Registro.data.desc()).all()

        registros_agrupados = {}
        for registro in all_registros:
            data_str = registro.data
            if data_str not in registros_agrupados:
                registros_agrupados[data_str] = []
            estudante = Estudante.query.get(registro.estudante_id)
            registros_agrupados[data_str].append(
                {'registro': registro, 'estudante_nome': estudante.nome if estudante else 'Desconhecido'})

        datas_ordenadas = sorted(registros_agrupados.keys(), reverse=True)

        return render_template('educador/ver_registros.html',
                               registros_agrupados=registros_agrupados,
                               datas_ordenadas=datas_ordenadas,
                               is_student_history=False)


@app.route('/excluir_registro/<int:registro_id>')
def excluir_registro(registro_id):
    """Exclui um registro de aula."""
    registro = Registro.query.get_or_404(registro_id)

    if not (registro.educador_id == session.get('usuario_id') or
            tipo_usuario() in ['Moderador', 'Direcao', 'Coordenacao']):
        flash('Você não tem permissão para excluir este registro.', 'danger')
        return redirect(url_for('painel_educador'))

    estudante_id_origem = registro.estudante_id
    db.session.delete(registro)
    db.session.commit()
    flash('Registro excluído com sucesso!', 'success')

    from_student_history = request.args.get('from_student_history', 'false').lower() == 'true'
    if from_student_history:
        return redirect(url_for('ver_registros', estudante_id=estudante_id_origem))
    else:
        return redirect(url_for('ver_registros'))


# Nova rota para gerar DOCX dos registros de um educador, agrupados por data
@app.route('/download_registros_educador_docx/<int:educador_id>')
def download_registros_educador_docx(educador_id):
    """Gera e baixa um documento DOCX com os registros de aula de um educador, agrupados por data."""
    if tipo_usuario() not in ['Moderador', 'Direcao', 'Coordenacao']:
        flash('Acesso não autorizado.', 'danger')
        return redirect(url_for('index'))

    educador = Usuario.query.get_or_404(educador_id)
    all_registros = Registro.query.filter_by(educador_id=educador.id).order_by(Registro.data.desc()).all()

    registros_agrupados = {}
    for registro in all_registros:
        data_str = registro.data
        if data_str not in registros_agrupados:
            registros_agrupados[data_str] = []
        estudante = Estudante.query.get(registro.estudante_id)
        registros_agrupados[data_str].append(
            {'registro': registro, 'estudante_nome': estudante.nome if estudante else 'Desconhecido'})

    datas_ordenadas = sorted(registros_agrupados.keys(), reverse=True)

    document = Document()
    document.add_heading(f'Registros de Aulas - Educador: {educador.nome_usuario}', level=1)
    document.add_paragraph(f'Gerado em: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
    document.add_paragraph('\n')  # Adiciona uma linha em branco

    if not all_registros:
        document.add_paragraph("Nenhum registro encontrado para este educador.")

    for data in datas_ordenadas:
        document.add_heading(f'Registros de {data}', level=2)

        for item in registros_agrupados[data]:
            r = item['registro']
            estudante_nome = item['estudante_nome']

            try:
                data_formatada = datetime.strptime(r.data, '%Y-%m-%d').strftime('%d/%m/%Y')
            except ValueError:
                data_formatada = r.data

            document.add_paragraph(f"Estudante: {sanitize_text_for_docx(estudante_nome)}", style='List Bullet')
            document.add_paragraph(f"Disciplina: {sanitize_text_for_docx(r.disciplina)}")
            document.add_paragraph(f"Quantidade de Aulas: {r.qtd_aulas}")
            document.add_paragraph("Relato:")
            document.add_paragraph(sanitize_text_for_docx(r.texto))
            document.add_paragraph('\n')  # Adiciona uma linha em branco entre os registros

    # Salva o documento em um buffer de memória
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        download_name=f"registros_educador_{educador.nome_usuario}.docx",
        as_attachment=True,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


# -------------------- INICIALIZAÇÃO MANUAL -------------------- #

def inicializar_banco():
    """Inicializa o banco de dados e cria um usuário moderador padrão se não existir."""
    with app.app_context():
        db.create_all()
        if not Usuario.query.filter_by(nome_usuario='SAVIO', tipo_acesso='Moderador').first():
            senha_hash = generate_password_hash('Ws396525$')
            moderador = Usuario(nome_usuario='SAVIO', senha_hash=senha_hash, tipo_acesso='Moderador', status='aprovado')
            db.session.add(moderador)
            db.session.commit()
            print("Usuário moderador 'SAVIO' criado com sucesso.")
        else:
            print("Usuário moderador 'SAVIO' já existe.")


if __name__ == '__main__':
    inicializar_banco()
    app.run(debug=True)
